#!/usr/bin/env python3
"""Generate RMA Management System User Guide (English) — Word Document"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Heading styles
for level in [1, 2, 3]:
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)  # dark green

# ── Cover / Title Page ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RMA Management System\nUser Guide')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('How to Submit & Batch Import After-Sales (RMA) Records')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'Version 3.6  |  July 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. System Overview',
    '2. Getting Started — Login & Navigation',
    '3. Method 1: Manually Create a New RMA Record',
    '   3.1  Open the New Record Form',
    '   3.2  Fill in Basic Information',
    '   3.3  Add RMA Items (Detail Lines)',
    '   3.4  Set Approvers',
    '   3.5  Submit for Approval or Save as Draft',
    '4. Method 2: Batch Import Historical Records',
    '   4.1  Download the Import Template',
    '   4.2  Fill in the Template',
    '   4.3  Upload and Preview',
    '   4.4  Review & Confirm Import',
    '   4.5  Submit Imported Drafts for Approval',
    '5. What Happens After Submission?',
    '6. Field Reference',
    '7. Frequently Asked Questions',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if not item.startswith('   '):
        for run in p.runs:
            run.font.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════
doc.add_heading('1. System Overview', level=1)

overview_text = (
    'The RMA (Return Merchandise Authorization) Management System is a web-based platform '
    'for tracking and managing after-sales service records. It supports:'
)
doc.add_paragraph(overview_text)

bullets = [
    'Manual entry of individual RMA cases with multiple SKU line items',
    'Batch import of historical records via Excel/CSV files',
    'Multi-level approval workflow (Level 1 → Level 2 → Level 3)',
    'Processing status tracking (ERP Receipt, Box Replacement, Parts Supply, RMA, Scrap)',
    'Export to Excel for factory analysis',
    'Lark (Feishu) notifications for approvers',
    'Role-based access control (Admin, Operator, Viewer)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ════════════════════════════════════════════════════════
doc.add_heading('2. Getting Started — Login & Navigation', level=1)

doc.add_paragraph(
    'Open your web browser and navigate to the system URL provided by your administrator. '
    'Sign in using your company Lark (Feishu) account, or use a local account if configured.'
)

doc.add_heading('Left Sidebar Navigation', level=2)
nav_items = [
    ('Overview', 'Dashboard with RMA statistics, return rates, and charts'),
    ('RMA Cases', 'View, search, filter, and manage all RMA records'),
    ('RMA Approvals', 'Review and approve/reject RMA cases assigned to you'),
    ('Products', 'Manage the product catalog (SKU database)'),
    ('Users', 'Manage system users (Admin only)'),
    ('Roles', 'Manage roles and permissions (Admin only)'),
    ('Approval Flow', 'Configure approval workflow rules (Admin only)'),
]
for name, desc in nav_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.font.bold = True
    p.add_run(desc)

# ════════════════════════════════════════════════════════
# 3. METHOD 1: MANUAL ENTRY
# ════════════════════════════════════════════════════════
doc.add_heading('3. Method 1: Manually Create a New RMA Record', level=1)
doc.add_paragraph(
    'Use this method when you need to enter a single RMA case — for example, when a customer '
    'reports a product issue and you want to log it immediately.'
)

# 3.1
doc.add_heading('3.1  Open the New Record Form', level=2)
doc.add_paragraph(
    '1. Click "RMA Cases" in the left sidebar.\n'
    '2. Click the green "Create Manually" button at the top of the list.\n'
    '3. A modal form titled "Create RMA Case" will appear.'
)

# 3.2
doc.add_heading('3.2  Fill in Basic Information', level=2)
doc.add_paragraph(
    'In the "Basic Information" section, fill in the required field:'
)
basic_table = doc.add_table(rows=2, cols=3)
basic_table.style = 'Light Grid Accent 1'
basic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = basic_table.rows[0].cells
hdr[0].text = 'Field'
hdr[1].text = 'Required?'
hdr[2].text = 'Description'
row = basic_table.rows[1].cells
row[0].text = 'RMA Date'
row[1].text = 'Yes'
row[2].text = 'The date when the after-sales issue occurred. Click the input and select from the calendar.'
for cell in basic_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
doc.add_paragraph()

# 3.3
doc.add_heading('3.3  Add RMA Items (Detail Lines)', level=2)
doc.add_paragraph(
    'Each RMA case can contain multiple items. Click the "+ Add Item" button to add rows. '
    'The following table explains each column:'
)

detail_table = doc.add_table(rows=8, cols=4)
detail_table.style = 'Light Grid Accent 1'
detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Column', 'Required?', 'Description', 'Example']
for i, h in enumerate(headers):
    detail_table.rows[0].cells[i].text = h
    for p in detail_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

data = [
    ['Platform', 'Yes', 'The sales platform where the return came from', 'Shopee, Lazada, TikTok Shop, Amazon, etc.'],
    ['SKU', 'Yes', 'Product SKU code. Type to search from the product database.', 'SKU-APPLE-001'],
    ['Order No.', 'No', 'Original order number for reference', 'SH20260115001'],
    ['Qty', 'Yes', 'Quantity returned. Must be > 0.', '2'],
    ['Return Reason', 'Yes', 'Reason for return. Select from the dropdown list.', 'Hardware Issue / Resellable / etc.'],
    ['Issue Description', 'Yes', 'Detailed description of the problem', 'Screen cracked, battery swollen'],
    ['Actions', '—', 'Click the × button to remove this row', '—'],
]
for i, row_data in enumerate(data):
    row = detail_table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
doc.add_paragraph()

doc.add_heading('Return Reason Options', level=3)
reason_table = doc.add_table(rows=6, cols=2)
reason_table.style = 'Light Grid Accent 1'
reason_table.alignment = WD_TABLE_ALIGNMENT.CENTER
reason_table.rows[0].cells[0].text = 'Chinese Value'
reason_table.rows[0].cells[1].text = 'English Meaning'
for p in reason_table.rows[0].cells[0].paragraphs:
    for r in p.runs: r.font.bold = True
for p in reason_table.rows[0].cells[1].paragraphs:
    for r in p.runs: r.font.bold = True

reasons = [
    ('可二次销售', 'Resellable — can be resold after return'),
    ('彩盒损坏', 'Damaged Color Box — packaging damage only'),
    ('配件缺失', 'Missing Parts — accessories missing'),
    ('硬件故障', 'Hardware Issue — functional defect'),
    ('报废', 'Scrap — cannot be repaired, dispose'),
]
for i, (cn, en) in enumerate(reasons):
    reason_table.rows[i+1].cells[0].text = cn
    reason_table.rows[i+1].cells[1].text = en
doc.add_paragraph()

# 3.4
doc.add_heading('3.4  Set Approvers', level=2)
doc.add_paragraph(
    'In the "Approver Settings" section, select at least one approver:\n'
    '• Level 1 Approver (optional)\n'
    '• Level 2 Approver (optional)\n'
    '• Level 3 Approver (optional)\n\n'
    'Important rules:\n'
    '• You must select at least ONE approver — the form will not submit with zero.\n'
    '• You can select the same person for multiple levels if needed.\n'
    '• The approval will flow in the order you filled (skipping empty levels).\n'
    '  - Example A: Only Level 1 filled → Level 1 approves → Done.\n'
    '  - Example B: Level 1 & Level 3 filled → Level 1 → Level 3 → Done.\n'
    '  - Example C: Only Level 2 filled → Level 2 approves → Done.'
)

# 3.5
doc.add_heading('3.5  Submit for Approval or Save as Draft', level=2)
doc.add_paragraph(
    'Two buttons at the bottom of the form:\n\n'
    '• "Submit for Approval" — Validates all fields and sends the record into the approval '
    'workflow. The first filled-in approver will receive a Lark notification.\n\n'
    '• "Save Draft" — Saves the record without triggering approval. You can return to edit '
    'and submit it later from the RMA Cases list. This is useful if you need to gather more '
    'information before submitting.\n\n'
    'Click "Cancel" to close the form without saving.'
)

doc.add_paragraph()

# Tip box (styled)
tip = doc.add_paragraph()
run = tip.add_run('💡 TIP: ')
run.font.bold = True
tip.add_run(
    'If you are entering a large number of historical records, use the Batch Import method '
    '(Section 4) instead of manual entry — it is much faster.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 4. BATCH IMPORT
# ════════════════════════════════════════════════════════
doc.add_heading('4. Method 2: Batch Import Historical Records', level=1)

doc.add_paragraph(
    'This is the recommended method for uploading past after-sales records in bulk. '
    'The system supports Excel (.xlsx) and CSV files.'
)

# 4.1
doc.add_heading('4.1  Download the Import Template', level=2)
doc.add_paragraph(
    '1. Go to "RMA Cases" from the left sidebar.\n'
    '2. Click the "Batch Import" button.\n'
    '3. In the dialog that opens, click "Download Import Template".\n'
    '4. An Excel file will be downloaded. It contains:\n'
    '   - Sheet 1 ("售后记录导入模板"): The data template with headers and an example row.\n'
    '   - Sheet 2 ("填写说明 Instructions"): Bilingual instructions and a reference table '
    'of valid return reasons in Chinese, English, and Bahasa Indonesia.'
)

# 4.2
doc.add_heading('4.2  Fill in the Template', level=2)
doc.add_paragraph(
    'The template has the following columns. Fill in one row per returned item:'
)

template_table = doc.add_table(rows=8, cols=3)
template_table.style = 'Light Grid Accent 1'
template_table.alignment = WD_TABLE_ALIGNMENT.CENTER
th = ['Column', 'Required?', 'How to Fill']
for i, h in enumerate(th):
    template_table.rows[0].cells[i].text = h
    for p in template_table.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True

tdata = [
    ['售后日期 (RMA Date)', 'Yes', 'Date in YYYY-MM-DD format, e.g. 2026-01-15'],
    ['平台 (Platform)', 'Yes', 'Platform name, e.g. Shopee, Lazada, Amazon'],
    ['SKU', 'Yes', 'Product SKU code exactly as in the product database'],
    ['订单号 (Order No.)', 'No', 'Original order number (leave blank if unknown)'],
    ['数量 (Qty)', 'Yes', 'A positive number, e.g. 1, 2, 3'],
    ['退货原因 (Return Reason)', 'Yes', 'Must match one of: 可二次销售, 彩盒损坏, 配件缺失, 硬件故障, 报废'],
    ['问题说明/备注 (Description)', 'No', 'Issue description or any notes'],
]
for i, row_data in enumerate(tdata):
    row = template_table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
doc.add_paragraph()

doc.add_heading('Important Notes for Template Filling', level=3)
notes = [
    'The column headers must be present in the first row. The system detects columns by header keywords (e.g. "日期", "platform", "sku", "qty", "退货原因"). If you use completely different headers, the import will fail.',
    'The "Return Reason" column accepts values in Chinese, English, or Bahasa Indonesia. The system will automatically map display names to internal values (e.g. "Hardware Issue" → "硬件故障").',
    'If an Excel file is used, the Return Reason column has a dropdown data validation list for easy selection.',
    'Leave blank rows empty — they will be skipped automatically.',
    'CSV files must be UTF-8 encoded. Excel files (.xlsx) are preferred because they handle special characters better.',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

# 4.3
doc.add_heading('4.3  Upload and Preview', level=2)
doc.add_paragraph(
    '1. In the Batch Import dialog, click the upload area or drag and drop your file.\n'
    '2. The system will parse the file and show a preview:\n'
    '   - How many rows were read\n'
    '   - How many rows are valid (will be imported)\n'
    '   - How many rows failed validation (with error reasons)\n'
    '3. Review any failed rows and fix the issues in your file. Re-upload if needed.'
)

# 4.4
doc.add_heading('4.4  Review & Confirm Import', level=2)
doc.add_paragraph(
    'After reviewing the preview, click "Import Valid Rows" to finish.\n\n'
    '⚠️ IMPORTANT: Imported records are saved as DRAFTS, not submitted for approval. '
    'This means:\n'
    '• They appear in the RMA Cases list with a "Draft" status badge.\n'
    '• No approval workflow is triggered yet.\n'
    '• No Lark notification is sent.\n'
    '• You (or the submitter) must manually edit each draft and click "Submit for Approval" '
    'to send it into the approval pipeline.'
)

# 4.5
doc.add_heading('4.5  Submit Imported Drafts for Approval', level=2)
doc.add_paragraph(
    'After batch import, complete these steps for each draft record:\n\n'
    '1. In the RMA Cases list, filter by status "Draft".\n'
    '2. Click on a draft record to open it.\n'
    '3. Review the data — fill in any missing fields if needed.\n'
    '4. Select at least one approver in the Approver Settings section.\n'
    '5. Click "Submit for Approval".\n\n'
    'Alternatively, you can delete incorrect drafts and re-import with corrected data.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 5. WHAT HAPPENS AFTER SUBMISSION
# ════════════════════════════════════════════════════════
doc.add_heading('5. What Happens After Submission?', level=1)

doc.add_heading('Approval Flow', level=2)
doc.add_paragraph(
    'Once submitted, the record enters the approval workflow:\n\n'
    '1. The first filled-in approver receives a Lark notification.\n'
    '2. They open the system, go to "RMA Approvals", review the case, and choose:\n'
    '   - Approve → moves to the next filled approver level (or marks "Approved" if this is the last one).\n'
    '   - Reject → record status changes to "Rejected", no further approval.\n'
    '3. After all filled-in approvers have approved, the record status becomes "Approved".\n'
    '4. Once approved, the processing workflow begins (see below).'
)

doc.add_heading('Processing Status Types', level=2)
doc.add_paragraph(
    'Each line item in a record is automatically assigned a processing type based on '
    'the return reason:'
)

proc_table = doc.add_table(rows=6, cols=2)
proc_table.style = 'Light Grid Accent 1'
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
proc_table.rows[0].cells[0].text = 'Return Reason'
proc_table.rows[0].cells[1].text = 'Auto-Assigned Processing Type'
for p in proc_table.rows[0].cells[0].paragraphs:
    for r in p.runs: r.font.bold = True
for p in proc_table.rows[0].cells[1].paragraphs:
    for r in p.runs: r.font.bold = True

procs = [
    ('Resellable (可二次销售)', 'Pending ERP Receipt'),
    ('Damaged Color Box (彩盒损坏)', 'Pending Box Replacement'),
    ('Missing Parts (配件缺失)', 'Pending Parts Supply'),
    ('Hardware Issue (硬件故障)', 'Pending RMA'),
    ('Scrap (报废)', 'Pending Disposal'),
]
for i, (reason, proc) in enumerate(procs):
    proc_table.rows[i+1].cells[0].text = reason
    proc_table.rows[i+1].cells[1].text = proc
doc.add_paragraph()

doc.add_paragraph(
    'Processing progress has three stages: Pending → Processing → Completed. '
    'Approved items can be updated by users with appropriate permissions through the '
    '"RMA Cases" list\'s batch update feature.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 6. FIELD REFERENCE
# ════════════════════════════════════════════════════════
doc.add_heading('6. Field Reference', level=1)

doc.add_heading('Aftersales Record (RMA Case)', level=2)
ref_table = doc.add_table(rows=18, cols=3)
ref_table.style = 'Light Grid Accent 1'
ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Field', 'Type', 'Description']):
    ref_table.rows[0].cells[i].text = h
    for p in ref_table.rows[0].cells[i].paragraphs:
        for r in p.runs: r.font.bold = True

ref_data = [
    ['RMA Date', 'Date (required)', 'Date the after-sales issue happened'],
    ['Platform', 'Text (required)', 'Sales platform name'],
    ['SKU', 'Text (required)', 'Product SKU code'],
    ['Order No.', 'Text (optional)', 'Original sales order number'],
    ['Qty', 'Number (required)', 'Quantity returned (> 0)'],
    ['Return Reason', 'Select (required)', 'One of 5 predefined reasons'],
    ['Issue Description', 'Text (required)', 'Detailed problem description'],
    ['Level 1 Approver', 'Select (optional)', 'First-level approver'],
    ['Level 2 Approver', 'Select (optional)', 'Second-level approver'],
    ['Level 3 Approver', 'Select (optional)', 'Third-level approver'],
    ['Approval Status', 'Auto', 'Draft → Pending L1 → L1 Approved → Pending L2 → ... → Approved / Rejected'],
    ['Processing Type', 'Auto', 'ERP / Box Replacement / Parts / RMA / Scrap'],
    ['Processing Progress', 'Manual', 'Pending / Processing / Completed'],
    ['Submitter', 'Auto', 'The user who created the record'],
    ['Created At', 'Auto', 'Timestamp of creation'],
    ['Updated At', 'Auto', 'Timestamp of last modification'],
    ['Brand', 'Auto', 'Derived from the SKU\'s product info'],
]
for i, row_data in enumerate(ref_data):
    row = ref_table.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 7. FAQ
# ════════════════════════════════════════════════════════
doc.add_heading('7. Frequently Asked Questions', level=1)

faqs = [
    (
        'Q: Can I submit a record without any approver?',
        'No. You must select at least one approver. The "Submit for Approval" button will show an error if no approver is selected.'
    ),
    (
        'Q: What if the SKU I entered does not exist in the product database?',
        'The record will still be saved, but the Brand, Model, and Category fields will be left blank. We recommend adding the product to the Product Management page first for complete data.'
    ),
    (
        'Q: I imported 100 records, but only 80 were valid. What happened?',
        'The preview dialog shows exactly which rows failed and why. Common causes: missing required fields, invalid return reason values, or negative quantities. Fix the file and re-import.'
    ),
    (
        'Q: How do I know which records are waiting for my approval?',
        'Go to "RMA Approvals" and select "Pending My Approval" from the scope dropdown. This shows only records assigned to you as the current approver.'
    ),
    (
        'Q: Can I change a record after submission?',
        'Only users with "Edit RMA Cases" permission can modify records. Contact your administrator if you need a record edited.'
    ),
    (
        'Q: What do the different approval statuses mean?',
        'Draft: not yet submitted. Pending Level 1/2/3 Approval: waiting for that level\'s approver. Level X Approved: that level\'s approver has approved, waiting for the next. Approved: all approvers have approved. Rejected: an approver rejected the case.'
    ),
    (
        'Q: Can the same person be all three levels of approver?',
        'Yes, the system allows the same person to be selected for multiple approval levels.'
    ),
    (
        'Q: Should I use Manual Entry or Batch Import?',
        'For ongoing day-to-day cases, use Manual Entry. For uploading many past/historical records at once, use Batch Import.'
    ),
    (
        'Q: My batch import file uses different column headers. Will it work?',
        'The system uses keyword matching on headers (e.g. it looks for words like "date", "platform", "sku", "qty", "return reason"). For best results, download and use the official template.'
    ),
]
for q, a in faqs:
    p_q = doc.add_paragraph()
    run = p_q.add_run(q)
    run.font.bold = True
    p_q.paragraph_format.space_after = Pt(2)

    p_a = doc.add_paragraph(a)
    p_a.paragraph_format.space_after = Pt(12)

# ── Footer ──
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    '— End of Guide —\n'
    'For technical support, please contact your system administrator.'
)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)

# ── Save ──
output_path = '/Users/a1-6/WorkBuddy/2026-07-01-18-44-52/aftersales-app/RMA_System_User_Guide_EN.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
